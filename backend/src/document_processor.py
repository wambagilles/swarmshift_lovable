"""
Module de traitement et de vectorisation des documents pour l'application Streamlit RAG No-Code
"""
import os
import tempfile
import uuid
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple

import PyPDF2
from docx import Document
import fitz  # PyMuPDF
from langchain_text_splitters import RecursiveCharacterTextSplitter, CharacterTextSplitter
#from langchain_community.vectorstores import Chroma
from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain_core.embeddings import Embeddings
from langchain_core.documents import Document as LangchainDocument

from chromadb import PersistentClient
from src.chroma_manager import chroma_manager


from src.config import DEV_MODE, TEMP_DIR, UPLOADS_DIR
from src.web_scraper import extract_text_from_url
from src.auth import get_current_user
from src.user_store import get_user_vector_store_path, get_collection_name, register_collection
from src.user_store_session import get_logged_user_id, get_session_value

# Fonction pour extraire le texte d'un fichier PDF avec numéros de page
def extract_text_from_pdf(file_path: str) -> List[Dict[str, Any]]:
    """
    Extrait le texte d'un fichier PDF avec les numéros de page
    
    Args:
        file_path: Chemin vers le fichier PDF
        
    Returns:
        Liste de dictionnaires contenant le texte et les métadonnées par page
    """
    result = []
    
    try:
        # Utiliser PyMuPDF (fitz) pour une extraction plus précise
        doc = fitz.open(file_path)
        
        for page_num, page in enumerate(doc):
            text = page.get_text()
            if text.strip():  # Ne pas inclure les pages vides
                result.append({
                    "content": text,
                    "page": page_num + 1,
                    "source": os.path.basename(file_path),
                    "path": file_path
                })
        
        doc.close()
    except Exception as e:
        # Fallback sur PyPDF2 si PyMuPDF échoue
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text and text.strip():  # Ne pas inclure les pages vides
                        result.append({
                            "content": text,
                            "page": page_num + 1,
                            "source": os.path.basename(file_path),
                            "path": file_path
                        })
        except Exception as inner_e:
            # Si les deux méthodes échouent, retourner une erreur
            result.append({
                "content": f"Erreur lors de l'extraction du texte: {str(e)} / {str(inner_e)}",
                "page": 0,
                "source": os.path.basename(file_path),
                "path": file_path,
                "error": True
            })
    
    return result

# Fonction pour extraire le texte d'un fichier Word avec numéros de page
def extract_text_from_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Extrait le texte d'un fichier Word avec estimation des numéros de page
    
    Args:
        file_path: Chemin vers le fichier Word
        
    Returns:
        Liste de dictionnaires contenant le texte et les métadonnées par paragraphe
    """
    result = []
    
    try:
        doc = Document(file_path)
        
        # Estimation: environ 3000 caractères par page
        chars_per_page = 3000
        current_chars = 0
        current_page = 1
        current_text = ""
        
        for para in doc.paragraphs:
            if para.text.strip():
                current_text += para.text + "\n"
                current_chars += len(para.text)
                
                # Si on dépasse l'estimation de caractères par page, on considère qu'on change de page
                if current_chars >= chars_per_page:
                    result.append({
                        "content": current_text,
                        "page": current_page,
                        "source": os.path.basename(file_path),
                        "path": file_path
                    })
                    current_page += 1
                    current_chars = 0
                    current_text = ""
        
        # Ajouter le reste du texte
        if current_text:
            result.append({
                "content": current_text,
                "page": current_page,
                "source": os.path.basename(file_path),
                "path": file_path
            })
    
    except Exception as e:
        result.append({
            "content": f"Erreur lors de l'extraction du texte: {str(e)}",
            "page": 0,
            "source": os.path.basename(file_path),
            "path": file_path,
            "error": True
        })
    
    return result

# Fonction pour découper le texte en chunks
def split_text(documents: List[Dict[str, Any]], 
               split_method: str = "recursive", 
               chunk_size: int = 1000, 
               chunk_overlap: int = 200) -> List[LangchainDocument]:
    """
    Découpe le texte en chunks selon la méthode spécifiée
    
    Args:
        documents: Liste de dictionnaires contenant le texte et les métadonnées
        split_method: Méthode de découpage ('recursive' ou 'character')
        chunk_size: Taille des chunks
        chunk_overlap: Chevauchement entre les chunks
        
    Returns:
        Liste de documents LangChain contenant les chunks et les métadonnées
    """
    langchain_docs = []
    
    # Sélectionner le splitter approprié
    if split_method == "recursive":
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    else:  # character
        splitter = CharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separator=" "
        )
    
    # Convertir les documents en format LangChain
    for doc in documents:
        if doc.get("error", False):
            # Si le document contient une erreur, l'ajouter tel quel sous forme de document LangChain
            langchain_docs.append(
                LangchainDocument(
                    page_content=doc["content"],
                    metadata={
                        "page": doc["page"],
                        "source": doc["source"],
                        "path": doc.get("path", ""),
                        "url": doc.get("url", ""),
                        "error": True
                    }
                )
            )
            continue
        
        # Créer un document LangChain
        langchain_doc = LangchainDocument(
            page_content=doc["content"],
            metadata={
                "page": doc["page"],
                "source": doc["source"],
                "path": doc.get("path", ""),
                "url": doc.get("url", "")
            }
        )
        
        # Ajouter le document à la liste pour le découpage
        langchain_docs.append(langchain_doc)
    
    # Découper les documents
    if langchain_docs:
        return splitter.split_documents(langchain_docs)
    
    return []

# Fonction pour obtenir la fonction d'embedding appropriée
def get_embedding_function(embedding_model: str = "openai") -> Embeddings:
    """
    Retourne la fonction d'embedding appropriée selon le modèle spécifié
    
    Args:
        embedding_model: Modèle d'embedding à utiliser ('openai' ou 'local')
        
    Returns:
        Fonction d'embedding LangChain
    """
    if embedding_model == "openai":
        # Vérifier que la clé API est disponible
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key and not DEV_MODE:
            raise ValueError("La clé API OpenAI n'est pas définie. Veuillez la définir dans le fichier .env")
        
        # En mode DEV, utiliser une clé factice si nécessaire
        if not api_key and DEV_MODE:
            api_key = "sk-dummy-key-for-dev-mode"
        
        return OpenAIEmbeddings(
            openai_api_key=api_key,
            model="text-embedding-ada-002"
        )
    else:
        # Utiliser un modèle local (HuggingFace)
        from langchain_community.embeddings import HuggingFaceEmbeddings
        
        return HuggingFaceEmbeddings(
            model_name="all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'}
        )

# Fonction pour vectoriser les documents avec Chroma
def vectorize_documents(documents: List[LangchainDocument], 
                        rag_id: str,
                        embedding_model: str = "openai") -> Tuple[str, Chroma]:
    """
    Vectorise les documents avec Chroma
    
    Args:
        documents: Liste de documents LangChain
        rag_id: ID du RAG pour lequel vectoriser les documents
        embedding_model: Modèle d'embedding à utiliser
        
    Returns:
        Tuple contenant le nom de la collection Chroma et l'instance Chroma
    """
    user_id = get_logged_user_id()
    # Obtenir la fonction d'embedding
    ef = get_embedding_function(embedding_model)
    
    # Obtenir le chemin de la base de données de l'utilisateur
    persist_directory = get_user_vector_store_path(user_id)
    
    # Obtenir le nom de la collection pour ce RAG
    
    collection_name = get_collection_name(user_id, rag_id)
    
    # Si aucune collection n'est associée à ce RAG, en créer une
    if not collection_name:
        collection_name = register_collection(rag_id)
    
    # Créer la base de données vectorielle
    
    chroma_client = chroma_manager.get_client()

    vector_store = Chroma.from_documents(
        collection_name=collection_name,
        embedding=ef,
        client=chroma_client,
        documents =documents
    )
    

    
    # Retrieve all documents from the collection
    docs = vector_store.get(include=["documents"])

    # Count the number of documents
    print(f"Nombre de documents dans la collection {collection_name}: {len(docs['documents'])}")

    #print(f"Nombre de documents dans la collection {collection_name}: {vector_store.__len__()}")
    
    return collection_name, vector_store

# Fonction principale pour traiter un document
def process_document(file_path: str, 
                     rag_id: str,
                     split_method: str = "recursive", 
                     chunk_size: int = 1000, 
                     chunk_overlap: int = 200,
                     embedding_model: str = "openai") -> Tuple[List[Dict[str, Any]], str, Chroma]:
    """
    Traite un document: extraction, découpage et vectorisation
    
    Args:
        file_path: Chemin vers le fichier
        rag_id: ID du RAG auquel le document est associé
        split_method: Méthode de découpage ('recursive' ou 'character')
        chunk_size: Taille des chunks
        chunk_overlap: Chevauchement entre les chunks
        embedding_model: Modèle d'embedding à utiliser
        
    Returns:
        Tuple contenant la liste des chunks (sous forme de dictionnaires), le nom de la collection Chroma,
        et l'instance Chroma
    """
    # Déterminer le type de fichier
    file_ext = os.path.splitext(file_path)[1].lower()
    
    # Extraire le texte selon le type de fichier
    if file_ext in ['.pdf']:
        documents = extract_text_from_pdf(file_path)
    elif file_ext in ['.docx', '.doc']:
        documents = extract_text_from_docx(file_path)
    else:
        # Fichier non supporté
        documents = [{
            "content": f"Type de fichier non supporté: {file_ext}",
            "page": 0,
            "source": os.path.basename(file_path),
            "path": file_path,
            "error": True
        }]
    
    # Découper le texte en chunks (documents LangChain)
    langchain_docs = split_text(documents, split_method, chunk_size, chunk_overlap)
    # Vectoriser les documents
    collection_name, vector_store = vectorize_documents(langchain_docs, rag_id, embedding_model)
    
    # Convertir les documents LangChain en dictionnaires pour la compatibilité avec le reste de l'application
    chunks = []
    for i, doc in enumerate(langchain_docs):
        chunks.append({
            "content": doc.page_content,
            "page": doc.metadata.get("page", 0),
            "source": doc.metadata.get("source", "inconnu"),
            "path": doc.metadata.get("path", ""),
            "url": doc.metadata.get("url", ""),
            "chunk_id": f"{doc.metadata.get('source', 'doc')}_{i}"
        })
    
    return chunks, collection_name, vector_store

# Fonction pour traiter une URL web
def process_web_url(url: str,
                    rag_id: str,
                    split_method: str = "recursive",
                    chunk_size: int = 1000,
                    chunk_overlap: int = 200,
                    embedding_model: str = "openai") -> Tuple[List[Dict[str, Any]], str, Chroma]:
    """
    Traite une URL web: extraction, découpage et vectorisation
    
    Args:
        url: URL de la page web
        rag_id: ID du RAG auquel le document est associé
        split_method: Méthode de découpage ('recursive' ou 'character')
        chunk_size: Taille des chunks
        chunk_overlap: Chevauchement entre les chunks
        embedding_model: Modèle d'embedding à utiliser
        
    Returns:
        Tuple contenant la liste des chunks (sous forme de dictionnaires), le nom de la collection Chroma,
        et l'instance Chroma
    """
    # Extraire le contenu de l'URL
    documents = extract_text_from_url(url)
    # Découper le texte en chunks (documents LangChain)
    langchain_docs = split_text(documents, split_method, chunk_size, chunk_overlap)
    # Vectoriser les documents
    collection_name, vector_store = vectorize_documents(langchain_docs, rag_id, embedding_model)
    # Convertir les documents LangChain en dictionnaires pour la compatibilité avec le reste de l'application
    chunks = []
    for i, doc in enumerate(langchain_docs):
        chunks.append({
            "content": doc.page_content,
            "page": doc.metadata.get("page", 0),
            "source": doc.metadata.get("source", "inconnu"),
            "path": doc.metadata.get("path", ""),
            "url": doc.metadata.get("url", ""),
            "chunk_id": f"{doc.metadata.get('source', 'web')}_{i}"
        })
    
    return chunks, collection_name, vector_store

# Fonction pour sauvegarder un fichier uploadé
def save_uploaded_file(uploaded_file) -> str:
    """
    Sauvegarde un fichier uploadé dans le répertoire temporaire
    
    Args:
        uploaded_file: Fichier uploadé via Streamlit
        
    Returns:
        Chemin vers le fichier sauvegardé
    """
    # Créer les répertoires si nécessaire
    os.makedirs(TEMP_DIR, exist_ok=True)
    os.makedirs(UPLOADS_DIR, exist_ok=True)
    
    # Générer un nom de fichier unique
    file_ext = os.path.splitext(uploaded_file.name)[1].lower()
    unique_filename = f"{uuid.uuid4().hex}{file_ext}"
    file_path = os.path.join(UPLOADS_DIR, unique_filename)
    
    # Sauvegarder le fichier
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    return file_path

# Fonction pour rechercher dans une collection Chroma
def search_documents(query: str, 
                     embedding_model: str = "openai",
                     n_results: int = 3) -> List[Dict[str, Any]]:
    """
    Recherche dans une collection Chroma
    
    Args:
        query: Requête de recherche
        rag_id: ID du RAG pour lequel effectuer la recherche
        embedding_model: Modèle d'embedding à utiliser
        n_results: Nombre de résultats à retourner
        
    Returns:
        Liste de dictionnaires contenant les résultats de recherche
    """
    user_id = get_logged_user_id()
    try:
        # Obtenir la fonction d'embedding
        ef = get_embedding_function(embedding_model)
        
        # Obtenir le chemin de la base de données de l'utilisateur
        persist_directory = get_user_vector_store_path(user_id)
        
        # Obtenir le nom de la collection pour ce RAG
        rag_id = get_session_value("current_rag")
        collection_name = get_collection_name(user_id, rag_id)
        
        if not collection_name:
            return [{
                "content": f"Erreur: Aucune collection trouvée pour le RAG {rag_id}",
                "page": 0,
                "source": "erreur",
                "path": "",
                "error": True
            }]
        
        # Charger la base de données vectorielle
        chroma_client = chroma_manager.get_client()
        vector_store = Chroma(
            collection_name=collection_name,
            embedding_function=ef,
            client=chroma_client
        )
        
        # Vérifier la taille de la collection
        
        # Retrieve all documents from the collection
        docs = vector_store.get(include=["documents"])

        # Count the number of documents
        collection_size = len(docs["documents"])

        print(f"Recherche dans la collection {collection_name} contenant {collection_size} documents")
        
        if collection_size == 0:
            return [{
                "content": f"Aucun document trouvé dans la collection {collection_name}",
                "page": 0,
                "source": "information",
                "path": "",
                "error": False
            }]
        
        # Effectuer la recherche
        results = vector_store.similarity_search(query, k=n_results)
        
        # Formater les résultats
        formatted_results = []
        
        for doc in results:
            formatted_results.append({
                "content": doc.page_content,
                "page": doc.metadata.get("page", 0),
                "source": doc.metadata.get("source", "inconnu"),
                "path": doc.metadata.get("path", ""),
                "url": doc.metadata.get("url", "")
            })
        
        return formatted_results
    
    except Exception as e:
        # En cas d'erreur, retourner un résultat vide
        print(f"Erreur lors de la recherche: {str(e)}")
        return [{
            "content": f"Erreur lors de la recherche: {str(e)}",
            "page": 0,
            "source": "erreur",
            "path": "",
            "error": True
        }]
